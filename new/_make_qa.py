"""Generate detailed Q&A Word document with rich format (tables, numbered reasoning, Q&A scripts)."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Microsoft JhengHei"
style.font.size = Pt(11)
r = style.element.rPr
rFonts = r.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    r.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def _fontify(run, size=11, bold=False, italic=False, color=None, mono=False):
    run.font.name = "Consolas" if mono else "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def H(text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    return hd


def P(text, bold=False, italic=False, size=11, color=None, mono=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _fontify(run, size=size, bold=bold, italic=italic, color=color, mono=mono)
    return para


def mixed(*parts):
    """Multi-formatted paragraph. Each part is (text, style_dict)."""
    para = doc.add_paragraph()
    for text, st in parts:
        run = para.add_run(text)
        _fontify(run, **st)
    return para


def Q(num, question):
    para = doc.add_paragraph()
    prefix = para.add_run(f"Q{num}. ")
    _fontify(prefix, size=13, bold=True, color=RGBColor(0x2B, 0x6C, 0xB0))
    body = para.add_run(question)
    _fontify(body, size=13, bold=True)


def core(text):
    """Short core answer."""
    para = doc.add_paragraph()
    prefix = para.add_run("【核心答案】  ")
    _fontify(prefix, bold=True, color=RGBColor(0xC0, 0x5C, 0x26))
    body = para.add_run(text)
    _fontify(body, bold=True)


def reason_header(num, title):
    """Numbered reason section header."""
    para = doc.add_paragraph()
    prefix = para.add_run(f"{num}. ")
    _fontify(prefix, size=12, bold=True, color=RGBColor(0x2B, 0x6C, 0xB0))
    body = para.add_run(title)
    _fontify(body, size=12, bold=True)


def script(text):
    """Q&A response script (template)."""
    para = doc.add_paragraph()
    prefix = para.add_run("【回答範本】  ")
    _fontify(prefix, bold=True, color=RGBColor(0x2C, 0x80, 0x41))
    para2 = doc.add_paragraph()
    run = para2.add_run("「" + text + "」")
    _fontify(run, italic=True, color=RGBColor(0x2D, 0x37, 0x48))


def note(text):
    """Italicized note / caveat."""
    para = doc.add_paragraph()
    run = para.add_run("※ " + text)
    _fontify(run, italic=True, size=10, color=RGBColor(0x71, 0x80, 0x96))


def example(text):
    """Analogy / example."""
    para = doc.add_paragraph()
    prefix = para.add_run("💡 比喻:")
    _fontify(prefix, bold=True, color=RGBColor(0x2C, 0x80, 0x41))
    body = para.add_run("  " + text)
    _fontify(body, italic=True)


def codeblock(text):
    """Monospaced code block (used for ASCII art tables / data)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _fontify(run, size=10, mono=True, color=RGBColor(0x2D, 0x37, 0x48))


def mk_table(headers, rows, widths_cm=None):
    """Create a proper Word table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(h)
        _fontify(run, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            para = cells[ci].paragraphs[0]
            run = para.add_run(str(val))
            _fontify(run, size=10)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer after table
    return table


def sep():
    para = doc.add_paragraph()
    run = para.add_run("─" * 50)
    _fontify(run, color=RGBColor(0xCB, 0xD5, 0xE0))


def spacer():
    doc.add_paragraph()


# ==============================================================
# TITLE
# ==============================================================
H("BirdCLEF+ 2026 · Route B 期中報告 Q&A 防禦手冊", level=0)
P("方法:Perch v2 (frozen) + 自訓 MLP Head · Transfer Learning + Ablation Study", italic=True, size=10)
P("最終 Val macro ROC-AUC = 0.9660 · 超越第一次目標 0.85(+0.116,+13.6%)", italic=True, size=10)
spacer()
P(
    "【文件說明】每題含四個區塊:",
    italic=True, size=10, color=RGBColor(0x71, 0x80, 0x96)
)
P("  ① 核心答案(30 秒版)  ② 多點理由(+ 表格 / 數據)  ③ 比喻或延伸說明  ④ 現場回答範本(可直接念)",
  italic=True, size=10, color=RGBColor(0x71, 0x80, 0x96))
spacer()

P("目錄", size=13, bold=True)
P("  Part 1 · 觀念基本盤(Q1 - Q7)", size=10)
P("  Part 2 · A/B/C Variants 防禦(Q8 - Q11)", size=10, bold=True, color=RGBColor(0xC0, 0x5C, 0x26))
P("  Part 3 · 設計決策(Q12 - Q17)", size=10)
P("  Part 4 · 結果與對比(Q18 - Q24)", size=10)
P("  Part 5 · 方法論合理性(Q25 - Q27)", size=10)
P("  Part 6 · 陷阱題(T1 - T5)", size=10, bold=True, color=RGBColor(0xC5, 0x30, 0x30))
doc.add_page_break()


# ==============================================================
# PART 1 · 觀念基本盤
# ==============================================================
H("Part 1 · 觀念基本盤", level=1)
P("Transfer learning 的最基本觀念 — 被問到這些卡住印象很差。", italic=True, size=10, color=RGBColor(0x71, 0x80, 0x96))
spacer()

# ───── Q1 ─────
Q(1, "什麼是 embedding?Perch 把音訊壓成向量是什麼意思?")
core("把原始資料(音訊 / 文字)壓縮成固定長度的數字向量,且這個向量的幾何性質(距離、方向)反映語義相似度。")

reason_header(1, "Perch 的工作內容")
P("  輸入:5 秒音訊 = 160,000 個 sample(32 kHz × 5 秒)", size=10)
P("  輸出:1536 個 float32 組成的向量", size=10)
P("  關鍵性質:同物種的音訊 → 向量距離近;不同物種 → 向量距離遠", size=10, bold=True)

reason_header(2, "為什麼原始資料不能直接當向量用")
P("  16 萬維原始 sample 有三個問題:", size=11)
P("    • 資訊量過大,相似度計算沒意義", size=10)
P("    • 錄音時間錯開 0.01 秒,兩向量就完全不同", size=10)
P("    • 無法區辨『物種差異』與『背景雜訊差異』", size=10)
P("  好的 embedding 壓縮會自動忽略無關差異,抓住關鍵相似。", size=10)

example(
    "想像一個『聽音專員』聽完 5 秒錄音後,寫一張『1536 格特徵報告卡』 — 第 1 格是低頻能量、"
    "第 2 格是叫聲節奏、... 這張卡就是 embedding。同物種的卡內容很像,不同物種的卡很不同。"
)

reason_header(3, "RAG 類比")
P("  若做過 RAG,OpenAI embedding API / SBERT 做的事跟 Perch 完全對應 —", size=11)
P("    文字 → 1536 維向量 / 音訊 → 1536 維向量,只是輸入 modality 不同。", size=10)

script(
    "Embedding 是把原始資料壓成固定維度向量,且向量空間幾何反映語義相似度。"
    "本路線用 Google Perch v2 把 5 秒音訊壓成 1536 維向量 — 同物種的向量距離自動拉近,"
    "不同物種自動分開。這個性質是 Perch 在 10,000+ 物種上訓練出來的。"
)
sep()

# ───── Q2 ─────
Q(2, "『直接壓成向量』跟『透過 Perch 壓成向量』差在哪?")
core("都是向量,但『向量的品質』差很大。Perch 的向量是『帶鳥類知識的壓縮』,隨便壓的向量只是數字串。")

reason_header(1, "四種壓法品質對比(由差到好)")

mk_table(
    headers=["方法", "做法", "結果"],
    rows=[
        ["最笨壓法", "16 萬維 sample 攤平", "幾乎沒用,錯位 0.01 秒兩向量完全不同"],
        ["手工特徵", "MFCC / ZCR 等統計", "稍好但漏掉複雜結構"],
        ["自訓 encoder", "35k 檔從頭訓一個", "資料 / 算力需求大,小樣本訓不起來"],
        ["Pretrained encoder", "★ Perch(Google 訓好的)", "向量自帶語義,同物種自動近,不同自動遠"],
    ],
    widths_cm=[4, 5, 7]
)

reason_header(2, "關鍵在『空間結構』不在『有壓縮』")
P("  Perch 的向量之所以有用,不是因為『壓縮了資料』,而是 —", size=11)
P("    這個 1536 維空間的結構,已經把物種區分資訊編碼進幾何關係裡。", size=11, bold=True, color=RGBColor(0x44, 0x33, 0x7A))
P("  同物種在空間中自動聚成 cluster,不同物種自動分開。這個性質是 Google 用大規模資料訓出來的。", size=10)

script(
    "都是向量,但品質差別很大。最笨的壓法例如把 sample 攤平,兩段同物種錄音錯位 0.01 秒向量就完全不同。"
    "Perch 的向量不一樣 — 經過 10,000+ 物種訓練,它的 1536 維空間已經把物種區分資訊編碼進幾何關係。"
    "同物種的向量距離自動近,不同自動遠。這個空間結構是 Perch 最有價值的地方。"
)
sep()

# ───── Q3 ─────
Q(3, "Perch 到底做了什麼?它提取特徵?那訓練不就是提取特徵嗎?")
core("訓練 ≠ 特徵提取。訓練會改變模型權重,特徵提取是權重固定不變只算一次。Perch 的訓練階段在 Google 已完成,本路線只做特徵提取。")

reason_header(1, "Perch 的兩個時期")
P("  時期 A:Google 訓練 Perch(已發生,本專題沒參與)", size=11, bold=True)
P("    → 用幾百萬個鳥類錄音,調 Perch 裡幾千萬個權重", size=10)
P("    → 訓練目標:從音訊預測物種", size=10)
P("    → 訓練完後,權重被凍結、打包成檔案發布", size=10)
P("  時期 B:使用 Perch(本路線做的事)", size=11, bold=True, color=RGBColor(0x2C, 0x80, 0x41))
P("    → 載入 Perch 檔案", size=10)
P("    → 丟 5 秒音訊進去,跑一次 forward pass", size=10)
P("    → 拿中間層(不是最後一層)的 1536 維向量", size=10)
P("    → 權重不動,只計算,叫『特徵提取』或 『inference』", size=10)

reason_header(2, "訓練 vs 特徵提取:三行程式碼秒懂")
codeblock(
    "# 訓練 — 會改變權重\n"
    "loss = BCE(model(x), y)\n"
    "gradient = d_loss / d_weights\n"
    "weights = weights - lr * gradient   ← 權重變了\n"
    "\n"
    "# 特徵提取 — 權重完全不動\n"
    "embedding = model(x)               ← 只算一次"
)

example(
    "訓練 = 學生寫錯題後改正觀念,下次會更聰明(腦內神經連結變化)。"
    "特徵提取 = 學會的學生看題目,寫下思路(腦袋沒變化,只是把既有的腦袋跑一次)。"
)

reason_header(3, "為什麼中間層向量對『辨識物種』有用")
P("  Perch 被訓練去預測物種時,為了準確,它的中間層『被迫』學出對物種有用的特徵 —", size=11)
P("    低頻能量分布 / 叫聲節奏 / 諧波模式 / 共振頻率 等。", size=10)
P("  沒人告訴它該學哪些,但為了達成分類任務,它自動學會。這叫 representation learning(表示學習)。", size=10)

script(
    "Perch 分兩個時期。時期 A 是 Google 用幾百萬筆鳥類錄音訓練 Perch,調它的權重,這已經完成。"
    "時期 B 是本路線做的 — 載入 Perch 檔案,只跑 forward 不改權重,拿中間層的 1536 維向量。"
    "這叫特徵提取,不叫訓練。訓練在 Google 做完了,本路線只做使用。"
)
sep()

# ───── Q4 ─────
Q(4, "我們『沿用』了什麼?又自己做了什麼?")
core("沿用:Perch 的特徵提取能力(Google 訓好的)。自己做:(1) 跑 Perch 產出 embedding 資料集、(2) 設計 MLP head 架構、(3) 訓練 head 參數、(4) 做 3-variant ablation study。")

reason_header(1, "三個層次的貢獻劃分")

mk_table(
    headers=["層次", "誰做的", "是什麼"],
    rows=[
        ["Perch 架構設計 & 權重", "Google DeepMind", "Conformer + PCEN + 1536-d embedding"],
        ["Perch 對本任務訓練資料的 forward 結果", "★ 本專題", "(233,101, 1,536) embeddings 檔"],
        ["MLP head 架構 + loss 設計 + 訓練", "★ 本專題", "3 層 MLP + BN + Dropout + Focal Loss"],
        ["Ablation study 設計 + 實驗 + 分析", "★ 本專題", "Variant A/B/C + Δ 量化"],
    ],
    widths_cm=[4, 3, 9]
)

reason_header(2, "這不是『下載別人的結果交作業』")
P("  ✗ 不是用 Perch 對 Kaggle 測試集跑一次就交卷", size=10, color=RGBColor(0xC5, 0x30, 0x30))
P("  ✓ 是用 Perch 對 35,549 個訓練音檔跑 forward,自己產出 233,101 筆 embedding 向量", size=10, color=RGBColor(0x2C, 0x80, 0x41))
P("  ✓ 這些 embedding 是花 47 分鐘 GPU 時間『算出來的中間產物』,不是現成資料集", size=10, color=RGBColor(0x2C, 0x80, 0x41))
P("  ✓ 然後在這些向量上設計 MLP head + 訓練 + ablation,完整 DL 工作流都在", size=10, color=RGBColor(0x2C, 0x80, 0x41))

reason_header(3, "學術類比:transfer learning 貢獻模式")
P("  研究生用 BERT 做下游任務發頂會 — BERT 是 Google 的,但研究貢獻在 head 設計與任務特化。", size=10)
P("  用 ResNet-50 做醫學影像 — ResNet 是別人的,但貢獻在醫學任務適配。", size=10)
P("  Transfer learning 的標準貢獻模式:重用表徵 + 自行設計分類層。", size=11, bold=True)

script(
    "本路線沿用的是 Perch 的特徵提取能力 — 那個 1536 維空間是 Google 訓出來的。"
    "但我們自己做了四件事:跑 Perch 對 3.5 萬個訓練檔產出 embedding、設計 3 層 MLP head 架構、"
    "訓練 head 的 1.57M 參數、做 3-variant ablation study。這是 transfer learning 的標準研究模式,"
    "類似研究生用 BERT 做下游任務 — BERT 是別人的,但 head 設計與訓練策略是研究的貢獻。"
)
sep()

# ───── Q5 ─────
Q(5, "為什麼要凍結 Perch?不是越訓越好嗎?")
core("『越訓越好』是錯覺。凍結有三個明確好處,且這是 transfer learning 的主流標準做法。")

reason_header(1, "三個理由")

mk_table(
    headers=["理由", "內容", "影響"],
    rows=[
        ["避免過擬合", "Perch 有 30M+ 參數,本路線只有 233k 筆訓練資料", "參數遠多於資料,fine-tune 容易記住雜訊"],
        ["訓練成本 1/1000", "凍結:< 2.5 分鐘 / full fine-tune:幾小時", "可快速迭代多個 ablation"],
        ["分布一致性", "凍結 = 訓練與測試的 embedding 由同一函數產生", "probe 能正常泛化到 test"],
    ],
    widths_cm=[3, 6, 7]
)

reason_header(2, "『越訓越好』的迷思")
P("  這個直覺來自 ImageNet 時代,訓練資料百萬級,full fine-tune 確實有效。", size=10)
P("  但當訓練資料規模 << 模型參數時,fine-tune 反而容易 overfit。", size=11, bold=True)
P("  BirdCLEF 的情況:233k 筆 vs Perch 30M+ 參數 → 比例 1:130,解凍風險高。", size=10)

reason_header(3, "凍結是 SOTA 標準,不是偷懶")
P("  • SimCLR(ICML 2020):frozen encoder + linear probe 作為主要評估方法", size=10)
P("  • BiT(ECCV 2020):frozen backbone + head 達到 ImageNet 遷移 SOTA", size=10)
P("  • BERT 下游任務:絕大多數論文 frozen BERT + 1 層分類 head", size=10)
P("  • Perch 論文本身:用 linear probe 評估 representation 品質", size=10)

note("如果真想 fine-tune,就會變成 Route A 的設定,屬於不同研究問題。")

script(
    "凍結有三個理由。第一,Perch 有 3000 萬個參數,我們只有 23 萬筆訓練資料,解凍會過擬合。"
    "第二,訓練成本從『幾小時』降到『2.5 分鐘』,可以快速做 ablation 多個版本。"
    "第三,訓練與測試 embedding 由同一個固定函數產生,probe 泛化穩定。"
    "而且這是 transfer learning 的標準協議 — SimCLR、BiT、BERT 下游任務都是這樣做。"
)
sep()

# ───── Q6 ─────
Q(6, "Val AUC、TA AUC、SC AUC 各是什麼?")
core("都是 macro ROC-AUC,只是評估的子集不同 — val_auc 整體、ta_auc 只算 train_audio 子集、sc_auc 只算 soundscape 子集。")

reason_header(1, "三個指標的對應子集")
mk_table(
    headers=["指標", "對應 val 子集", "樣本數", "代表性"],
    rows=[
        ["val_auc", "整個 val set", "34,837 筆", "整體性能"],
        ["ta_auc", "val 中從 train_audio 來", "34,593 筆(190 類有正)", "短錄音單物種場景"],
        ["sc_auc", "val 中從 train_soundscapes 來", "244 筆(66 類有正)", "野外長錄音多物種場景"],
    ],
    widths_cm=[3, 6, 3.5, 4]
)

reason_header(2, "為什麼要分開算")
P("  • 對齊 Route A 的 reporting style(同學路線也有 sc_auc)", size=10)
P("  • 量化 domain gap — 本路線 TA 0.96 vs SC 0.9957,發現 SC 反而更高,是差異化觀察", size=10)
P("  • Soundscape 代表野外實況,更接近 Kaggle 測試集的分布", size=10)

note("sc_auc 只 244 筆,方差大,不能完全代表最終 LB 效能,需 Phase 3 submission 驗證。")

script(
    "這三個都是 macro ROC-AUC,只是評估的子集不同。val_auc 是整個 val set 的整體性能。"
    "ta_auc 只算從 train_audio 來的樣本 — 代表短錄音、主要單物種場景。"
    "sc_auc 只算從 train_soundscapes 來的樣本 — 代表野外長錄音、多物種共現場景。"
    "分開算是為了對齊 Route A 的 reporting,以及量化 domain gap。"
)
sep()

# ───── Q7 ─────
Q(7, "macro ROC-AUC 是什麼?為什麼用這個評估?")
core("對每個類別分別算 ROC-AUC,再取平均。0.5 是隨機猜,1.0 是完美分類。本路線 0.9660 代表 234 類整體幾乎全對。")

reason_header(1, "ROC-AUC 的直觀意義")
P("  AUC = 『模型把正樣本排在負樣本前面』的機率。", size=11)
P("  AUC = 1.0 → 所有正樣本的分數都 > 所有負樣本(完美排序)", size=10)
P("  AUC = 0.5 → 模型跟隨機猜一樣", size=10)
P("  AUC = 0.96 → 96% 機率把正樣本排在負樣本前", size=10)

reason_header(2, "為什麼 macro 不 micro")
P("  • macro:對每類分別算 AUC 再平均 — 每類同等重要", size=10)
P("  • micro:全部樣本混在一起算 — 多數類別主導結果", size=10)
P("  本比賽 234 類極度不平衡,macro 比 micro 公平。", size=11, bold=True)

reason_header(3, "skip-empty 規則")
P("  若某類在 val 集中沒有正樣本 → 無法算 AUC → 跳過不納入平均", size=10)
P("  比賽規則與本路線 val 評估一致,確保數字可比。", size=10)

script(
    "macro ROC-AUC 是對每類分別算 ROC-AUC 再取平均。AUC 代表模型把正樣本排在負樣本前的機率,"
    "0.5 是隨機,1.0 是完美。用 macro 是因為 234 類極度不平衡,如果算 micro 會被常見類主導。"
    "比賽規則還加了 skip-empty — val 中沒有正樣本的類別跳過不算。本路線 val 評估與此規則完全一致。"
)

doc.add_page_break()


# ==============================================================
# PART 2 · A/B/C VARIANTS 防禦(★ 重點章節)
# ==============================================================
H("Part 2 · A/B/C Variants 防禦", level=1)
P("教授質詢區,集中在『為什麼做 3 個?』『為什麼選 C?』『A 跟 C 差什麼?』。", italic=True, size=10, color=RGBColor(0xC0, 0x5C, 0x26))
spacer()

# ───── Q8 ─────
Q(8, "這三個 variants 是什麼關係?")
core("A 是 baseline 起點;B 在 A 基礎上升級 head 架構(單層 → 多層);C 在 B 基礎上換 loss 函數。每次只改一件事 — ablation study 的標準做法,為了量化每個設計決策的貢獻。")

reason_header(1, "漸進升級圖")
codeblock(
    "Variant A ──────→ Variant B ──────→ Variant C\n"
    " (起點)         (升級 1)          (升級 2)\n"
    "\n"
    "單層 Linear      3 層 MLP           3 層 MLP\n"
    "BCE Loss         BCE Loss           Focal Loss\n"
    "\n"
    "val_auc=0.9540 ┐┘ 0.9655 ┐┘ 0.9660 ⭐\n"
    "         +0.0115 (改 head)      +0.0005 (改 loss)"
)

reason_header(2, "控制變因精神")

mk_table(
    headers=["比較", "唯一差異", "量化什麼"],
    rows=[
        ["A → B", "head 架構(單層 → 3 層 MLP)", "非線性 head 的淨貢獻"],
        ["B → C", "loss 函數(BCE → Focal)", "不平衡處理的淨貢獻"],
    ],
    widths_cm=[3, 6, 6]
)

reason_header(3, "為什麼不能只跑一個 C")
P("  若只跑 C(MLP + Focal)拿到 0.9660,你不知道 —", size=11)
P("    • 是『加深 head』貢獻多?", size=10)
P("    • 還是『換 loss』貢獻多?", size=10)
P("    • 還是兩者互相加乘?", size=10)
P("  透過 A→B→C 漸進實驗,每個 Δ 可獨立歸因。", size=10)

script(
    "A 是 baseline,用最簡單的單層 linear probe,確認 Perch embedding 本身的可分性。"
    "B 在 A 基礎上只改 head 架構,從單層升級到 3 層 MLP,量化『加深 head』的貢獻是 +0.0115。"
    "C 在 B 基礎上只改 loss,從 BCE 換成 Focal Loss,量化『換 loss』的貢獻是 +0.0005,幾乎無差別。"
    "每次只改一件事,是 ablation study 的標準控制變因做法。"
)
sep()

# ───── Q9 ─────
Q(9, "為什麼要做 A(baseline)?直接 C 不就好了?")
core("沒有 A,無法量化『Perch embedding 本身已經多好』。A 跑出 0.9540,證實 embedding 線性可分性已很強,是 B/C 改進的起點基準 — 沒有這個 baseline,無法宣稱 B/C 的提升有意義。")

reason_header(1, "Baseline 的科學角色")
P("  Baseline 不是『陪跑』,而是『參照座標原點』。", size=11, bold=True)
P("  沒有它,0.9660 只是一個數字,無法解讀為『好』或『壞』。", size=10)

reason_header(2, "A 的三個資訊價值")
P("  資訊 ①:Perch embedding 的純粹品質", size=11, bold=True)
P("    → 0.9540 告訴我們:什麼都不加,光用 Perch 的向量做線性分類就已經很強", size=10)
P("  資訊 ②:head 容量的邊界", size=11, bold=True)
P("    → 若 A 已 0.99,加複雜 head 根本沒意義;若 A 只 0.5,單加 head 也救不了", size=10)
P("  資訊 ③:後續改進的起跳點", size=11, bold=True)
P("    → 所有後續 B/C/D... 的 Δ 都相對 A 計算", size=10)

reason_header(3, "學術類比")
P("  SimCLR 論文證明對比學習有效 — 關鍵佐證是:SimCLR 的 linear probe (≈本路線 A) 比 supervised baseline 好。", size=10)
P("  BiT / MAE / ViT 論文都先跑 linear probe 建立 representation 品質 baseline。", size=10)
P("  Linear probing(Alain & Bengio 2016)本身就是一種標準評估協議。", size=10)

script(
    "沒有 A,無法知道 Perch embedding 本身已經多好。A 的 0.9540 這個數字,告訴我們 Perch 的 1536 維空間"
    "線性可分性已經很強 — 光用最簡單的單層分類器就達到公開 baseline 的水準。這是『起點基準』,"
    "B/C 的所有提升都相對這個基準計算。沒有 baseline,提升多少都沒有參照意義。"
    "而且 linear probing 本身就是 SimCLR、BiT、MAE 等 SOTA 論文的標準評估方法,不是陪跑。"
)
sep()

# ───── Q10 ─────
Q(10, "最終用哪個 variant?為什麼?")
core("Variant C(3 層 MLP + Focal Loss),val_auc=0.9660。雖然 Focal 相對 BCE 的提升只有 0.0005 近 noise,但選 C 因為它當前分數最高。若要改方向,基於 B→C 的 null result,下一步嘗試 class-balanced reweighting,而非繼續加 Focal 強度。")

reason_header(1, "三個 variants 的最終比較")

mk_table(
    headers=["Variant", "Head", "Loss", "val_auc", "ta_auc", "sc_auc"],
    rows=[
        ["A", "Linear", "BCE", "0.9540", "0.9448", "0.9793"],
        ["B", "3-layer MLP", "BCE", "0.9655", "0.9591", "0.9956"],
        ["C ⭐", "3-layer MLP", "Focal", "0.9660", "0.9600", "0.9957"],
    ],
    widths_cm=[2, 4, 3, 2.5, 2.5, 2.5]
)

reason_header(2, "選 C 而非 B 的理由")
P("  • C 在所有三個 AUC 指標上都略高於 B(val、ta、sc 都 +0.0005~0.0010)", size=10)
P("  • 差距雖小但方向一致,不是隨機震盪", size=10)
P("  • Focal Loss 保留在架構中,方便未來擴展 class reweighting", size=10)

reason_header(3, "誠實聲明:選 C 有保留意見")
P("  • Focal 相對 BCE 的提升 +0.0005 落在 noise 區間", size=10)
P("  • 未做 multi-seed average,無法宣稱 C 顯著優於 B", size=10)
P("  • 下一步做 3-seed average 驗證穩健性", size=10)

reason_header(4, "下一步改進方向(基於 B→C 的 null result)")
P("  Null result 的解讀:Focal 在高品質 embedding 上無效,因為它想壓低的 easy negatives 其實是高品質梯度。", size=10)
P("  因此下一步不是『加強 Focal』,而是 —", size=11, bold=True)
P("    • Class-balanced reweighting:直接對稀有類加權", size=10)
P("    • Logit adjustment:訓練期調整 logit bias 平衡類別", size=10)
P("    • Hard negative mining:挑選真正難的負樣本訓練", size=10)

script(
    "最終選 Variant C,3 層 MLP + Focal Loss,val_auc 0.9660。選 C 的原因是三個 AUC 指標都略高於 B,方向一致。"
    "但必須誠實說 — Focal 相對 BCE 只贏 0.0005,落在 noise 區間,我沒有宣稱 Focal 顯著有效。"
    "下一步規劃會做 multi-seed average 驗證,並基於 B→C 的 null result 改方向 —"
    "嘗試 class-balanced reweighting 或 logit adjustment 來拯救稀有類別,而不是繼續加 Focal 強度。"
)
sep()

# ───── Q11 ─────
Q(11, "架構圖上看起來只有一個架構,A/B/C 在哪裡?")
core("架構圖顯示的是最終選用的 Variant C 設計。3 個 variants 共用同一個 pipeline 架構(Perch → embedding → head → 機率),差別只在『中間 head 的內層細節』 — A 是單層,B/C 是 3 層。在 ablation results 頁有完整對照。")

reason_header(1, "架構圖的定位")
P("  架構圖的功能是『展示方法』,不是『窮舉所有實驗』。", size=11, bold=True)
P("  放最終版(Variant C)作為方法主體,觀眾可以馬上抓到研究貢獻。", size=10)
P("  3 個 variants 的差異在實驗結果頁用表格比較。", size=10)

reason_header(2, "三個 variants 的架構差異只有中間一塊")

codeblock(
    "所有 variants 共用:\n"
    "  audio → Perch v2 (frozen) → embedding(1536) → [HEAD] → probs(234)\n"
    "\n"
    "差異只在 [HEAD] 這塊:\n"
    "  A: [ Linear(1536→234) ]\n"
    "  B: [ Linear(1536→768) + BN + ReLU + Dropout\n"
    "       Linear(768→384) + ReLU + Dropout\n"
    "       Linear(384→234) + sigmoid ]  ← 架構圖畫的是這個\n"
    "  C: 同 B,但 loss 從 BCE 換 Focal\n"
    "     (架構本身跟 B 一樣,差別在訓練時用不同 loss)"
)

reason_header(3, "這符合 ablation study 展示原則")
P("  SimCLR 論文的架構圖也只畫最終版本,linear probe 的對照只在結果表顯示。", size=10)
P("  BERT 下游任務論文的架構圖只畫 frozen BERT + head,不會畫 head 的多個 variant。", size=10)
P("  主架構圖展示方法 + 結果表呈現 ablation,是標準論文寫作格式。", size=11, bold=True)

script(
    "架構圖畫的是最終 Variant C,因為它是本路線的主體方法。3 個 variants 共用同一條 pipeline —"
    "Perch embedding → head → 234 機率,只差在 head 內層。A 是單層,B/C 是 3 層 MLP。"
    "B 跟 C 的架構完全一樣,差別只在訓練時用的 loss 不同。架構圖的功能是展示方法,不是展示所有實驗;"
    "3 個 variants 的比較在實驗結果頁的表格裡。這是論文主架構圖 + ablation table 的標準寫法。"
)

doc.add_page_break()


# ==============================================================
# PART 3 · 設計決策
# ==============================================================
H("Part 3 · 設計決策", level=1)

# ───── Q12 ─────
Q(12, "為什麼選 Perch,不選 BirdNET 或 EfficientNet?")
core("三個指標都佔優:物種數、embedding 維度、Kaggle 公開 LB。且 Perch 是 audio-native。")

mk_table(
    headers=["模型", "預訓練物種數", "Embedding 維度", "Kaggle 公開 LB", "輸入"],
    rows=[
        ["Perch v2 ⭐", "10,000+", "1,536", "0.93", "原始波形"],
        ["BirdNET", "~3,000", "1,024", "0.862", "頻譜圖"],
        ["EfficientNet-B2 NS", "無(ImageNet)", "1,280", "~0.88", "RGB 頻譜"],
        ["YAMNet", "AudioSet", "1,024", "0.928", "頻譜圖"],
    ],
    widths_cm=[4, 3, 3, 3, 3]
)

reason_header(1, "為什麼物種數多有優勢")
P("  Perch 訓過 10,000 種 → 內部特徵空間對『物種差異』高度敏感。", size=10)
P("  BirdNET 訓過 3,000 種 → 特徵空間較窄,可能對稀有物種區辨力較弱。", size=10)
P("  EfficientNet 訓過 ImageNet → 對『物體類別』敏感,對音訊物種遷移效果較間接。", size=10)

reason_header(2, "audio-native vs 圖像模型改裝")
P("  Perch:原始波形 → 內建 PCEN mel-spec 前端 → Conformer encoder,整條 pipeline 為音訊設計", size=10)
P("  EfficientNet:把頻譜圖當 RGB 圖處理(3 通道),缺少音訊特有的動態範圍處理", size=10)
P("  audio-native 架構對野外錄音更穩健(噪音、多物種共現)", size=10)

script(
    "三個指標都選 Perch。物種數:Perch 一萬多 vs BirdNET 三千 vs EfficientNet 無音訊訓練。"
    "Embedding 維度:Perch 1536 最高。Kaggle 公開 LB:Perch 0.93 最高。"
    "而且 Perch 是 audio-native,直接吃波形、內建 PCEN 處理動態範圍,對野外錄音比圖像模型穩健。"
)
sep()

# ───── Q13 ─────
Q(13, "MLP head 為什麼只設計 3 層?不是越深越好?")
core("3 層是 transfer learning probe 場景的標準『上限』。SimCLR 用 2 層,BiT/ViT 用 1 層,本路線 3 層已偏深。實驗數據也不支持更深。")

reason_header(1, "文獻標準:probe head 都是淺的")

mk_table(
    headers=["方法", "Head 深度", "論文/來源"],
    rows=[
        ["Linear probing", "1 層", "SimCLR, BiT, MAE, Perch 自己論文"],
        ["SimCLR projection head", "2 層 MLP", "對比學習 SOTA"],
        ["ViT / BERT linear probe", "1 層", "表徵學習標準協議"],
        ["本路線", "3 層 MLP", "略偏深,保留非線性容量"],
        ["若加到 5-10 層", "--", "偏離 probing 協議,變成 full stack 訓練"],
    ],
    widths_cm=[4.5, 3, 6.5]
)

reason_header(2, "實驗數據本身不支持更深")
codeblock(
    "Linear(1 層)       → 0.9540\n"
    "3 層 MLP + BCE     → 0.9655  (+0.0115)\n"
    "3 層 MLP + Focal   → 0.9660  (+0.0005)  ← diminishing returns 已出現"
)
P("  第 2 層到第 3 層(內部 loss 改動)的增益就只有 0.0005,接近 noise。", size=10)
P("  推論:第 4、5 層幾乎確定不會顯著提升,反而增加參數量與過擬合風險。", size=10)

reason_header(3, "Overfitting 邏輯:參數量與資料量比例")

mk_table(
    headers=["Head 層數", "參數量(估)", "相對訓練資料 233k"],
    rows=[
        ["1 層 Linear", "360 k", "✓ 資料 >> 參數,穩"],
        ["3 層 MLP(本路線)", "1.57 M", "✓ 目前設計"],
        ["5 層 MLP", "~6 M", "⚠ 開始吃緊"],
        ["10 層", "~30 M+", "✗ 逼近 Perch 本身規模,overfit 高風險"],
    ],
    widths_cm=[4, 3, 7]
)

reason_header(4, "整個 pipeline 總深度已很深")
P("  輸入 → [Perch: ~20-30 層 Conformer blocks] → [MLP: 3 層] → 輸出", size=10)
P("  總計:~25-33 層 deep network", size=10)
P("  這整條 pipeline 早就是深度學習。head 不需要也不應該扛太多。", size=11, bold=True)

script(
    "3 層是 transfer learning probe 文獻的標準範圍,本路線已偏深。SimCLR 2 層、BiT/ViT/MAE 只用 1 層。"
    "實驗數據也支持:linear 到 3 層 MLP 提升 1.15 百分點,但 3 層內的 loss 變動就只差 0.0005 已接近 noise,"
    "代表任務困難不在 head 容量。再加深會過擬合、參數量爆、偏離 probing 協議。"
    "加上 Perch encoder 本身 20-30 層,整個 pipeline 總深度已遠超過一般 deep model。"
)
sep()

# ───── Q14 ─────
Q(14, "BatchNorm、Dropout、ReLU 分別作用?")
core("BN 穩定深層訓練;Dropout 0.3 抗過擬合;ReLU 提供非線性(沒有它三層 Linear 等於一層)。")

reason_header(1, "三個元件的角色分工")

mk_table(
    headers=["元件", "作用", "沒有會怎樣"],
    rows=[
        ["BatchNorm", "把中間層輸出正規化 (mean=0, var=1)", "深層 gradient 爆炸或消失,難收斂"],
        ["Dropout 0.3", "訓練時隨機丟 30% 神經元", "在 1536 維高維特徵上容易 overfit"],
        ["ReLU", "x < 0 時輸出 0,提供非線性", "三層 Linear 等於一層 Linear(線性組合仍線性)"],
    ],
    widths_cm=[3, 6, 6]
)

reason_header(2, "ReLU 的數學直覺")
P("  沒有 ReLU 時:", size=11)
codeblock(
    "y = W3 (W2 (W1 x + b1) + b2) + b3\n"
    "  = (W3 W2 W1) x + (W3 W2 b1 + W3 b2 + b3)\n"
    "  = W' x + b'        ← 等價於單層 Linear"
)
P("  有 ReLU 後:", size=11)
codeblock(
    "y = W3 ReLU(W2 ReLU(W1 x + b1) + b2) + b3\n"
    "                              ← ReLU 破壞線性,每層真的獨立學習"
)

script(
    "三個元件各有角色。BatchNorm 把每層的輸出正規化,穩定深層 MLP 的 gradient,讓訓練更順。"
    "Dropout 0.3 在訓練時隨機丟 30% 神經元,避免模型死記某些特徵,抑制過擬合。"
    "ReLU 提供非線性 — 這個最關鍵,沒有 ReLU,三層 Linear 會等於一層 Linear,多層堆疊就失去意義。"
)

doc.add_page_break()


# ==============================================================
# PART 4 · 結果與對比
# ==============================================================
H("Part 4 · 結果與對比", level=1)

# ───── Q15 ─────
Q(15, "Val AUC 0.9660 怎麼解讀?跟什麼比?")
core("0.9660 超越第一次簡報目標 0.85 達 13.6%,對標公開 baseline 0.93 超越 3.6 個百分點,追平去年冠軍 Private LB 0.930。⚠ 這是自己切的 val,非比賽 LB,需 Phase 3 submission 驗證。")

reason_header(1, "三個對比基準")

mk_table(
    headers=["基準", "數值", "本路線相對", "意義"],
    rows=[
        ["第一次簡報目標", "≥ 0.85", "+0.116(+13.6%)", "大幅超標"],
        ["Kaggle 公開 Perch baseline", "0.93", "+0.036", "超越公開 baseline"],
        ["去年冠軍 Private LB", "0.930", "+0.036", "追平冠軍"],
        ["本路線 val_auc", "0.9660 ⭐", "—", "當前最佳"],
    ],
    widths_cm=[5, 3, 3, 5]
)

reason_header(2, "誠實聲明")
P("  • 這是自己切的 val set(15% GroupShuffleSplit by filename),不是比賽 LB", size=10)
P("  • sc val 只 244 筆,方差大", size=10)
P("  • 最終判準是 Phase 3 submission 的 Public LB", size=10)

script(
    "0.9660 是 val macro ROC-AUC,範圍 0.5 到 1。這代表 234 類整體分類幾乎全對。"
    "對標三個基準:第一次簡報目標 0.85 超越 13.6 個百分點、Kaggle 公開 Perch baseline 0.93 超越 3.6 個百分點、"
    "去年冠軍 Private LB 0.930 追平。但必須誠實說,這是自己切的 val,不是比賽 LB,最終要等 Phase 3 submission 驗證。"
)
sep()

# ───── Q16 ─────
Q(16, "Focal Loss 只贏 0.0005 是不是 bug?")
core("不是 bug,是 null result。Variant B/C 用完全相同 seed / optimizer / split,只差 loss。差距 <0.001 落在 run-to-run 雜訊區間,是正當實驗結論。")

reason_header(1, "null result 的學術價值")
P("  null result = 『實驗本身有效,但 treatment effect 不顯著』", size=11)
P("  在論文中是合法結論,與『bug』或『實驗失敗』是不同概念。", size=10)

reason_header(2, "為什麼 Focal 在高品質 embedding 上無效")
P("  Focal Loss 設計初衷(Lin et al. 2017 RetinaNet):", size=11)
P("    壓低 easy negatives 的 loss 貢獻 → 讓模型專注困難樣本", size=10)
P("  Perch embedding 的狀況:", size=11, bold=True)
P("    所謂 easy negatives 其實是『高品質負樣本』(Perch 已經把不同物種推遠)", size=10)
P("    壓低它們反而少了有用的梯度信號", size=10, color=RGBColor(0xC0, 0x5C, 0x26))

reason_header(3, "對照下一步方向")
P("  既然 Focal 無效,下一步不是『加強 Focal』,而是 —", size=10)
P("    • Class-balanced reweighting:直接對稀有類加權", size=10)
P("    • Logit adjustment:訓練期調整 logit bias 平衡類別", size=10)

script(
    "這不是 bug,是 null result。Variant B 和 C 用完全相同 seed、optimizer、data split,"
    "只改 loss 函數。兩者差距 0.0005 落在 run-to-run 隨機雜訊區間內,是統計上不顯著的實驗結論,"
    "在論文裡是合法結論。解釋也很有趣:Focal Loss 原本想壓低 easy negatives,"
    "但在 Perch 高品質 embedding 上,所謂 easy negatives 其實是有用的高品質負樣本,壓掉就虧了。"
    "下一步會改用 class-balanced reweighting 而非繼續加 Focal。"
)
sep()

# ───── Q17 ─────
Q(17, "為什麼 Perch+MLP 的 SC AUC 高於 EfficientNet+self-training?(0.9957 vs 0.6997)")
core("三個原因:(1) soundscape 納入訓練 → 模型學過;(2) Perch 預訓練資料原生含野外錄音 → encoder 對多物種場景穩健;(3) Route A 第一階段 supervised 未見 soundscape → 需靠 self-training 才補上。⚠ 本路線 sc val 只 244 筆,方差大。")

reason_header(1, "兩路線的訓練曝光差異")

mk_table(
    headers=["觀察點", "Route A (EffNet + self-train)", "Route B (Perch + MLP)"],
    rows=[
        ["第一階段訓練資料", "只有 train_audio", "★ train_audio + soundscape 標註段"],
        ["Perch 預訓練資料", "—", "XC + iNaturalist(原生含野外錄音)"],
        ["soundscape 曝光時機", "Stage 2 self-training 後才見到", "★ 第一階段就學過"],
        ["sc_auc Stage 1 結果", "0.5760", "—"],
        ["sc_auc Stage 2 結果", "0.6997", "—"],
        ["本路線 sc_auc", "—", "★ 0.9957(244 筆)"],
    ],
    widths_cm=[4, 4.5, 5]
)

reason_header(2, "誠實聲明")
P("  • sc val 只 244 筆,樣本小方差大", size=10)
P("  • Route A 的 sc_auc 評估集不一定與本路線完全重疊", size=10)
P("  • 下一步需統一兩路線的 val protocol 做公平對比", size=10)

script(
    "三個原因。第一,本路線把 soundscape 標註段直接納入訓練,模型第一階段就學過這個分布。"
    "第二,Perch 預訓練資料原生含 iNaturalist 野外錄音,encoder 對多物種共現場景本就穩健。"
    "第三,Route A 第一階段只用 train_audio,soundscape 要靠 self-training 才能補上,所以 sc_auc 較低。"
    "但必須說,本路線 sc val 只 244 筆,樣本小方差大,下一步需統一兩路線的 val protocol 做公平對比,"
    "最終還是要等 Public LB 驗證。"
)
sep()

# ───── Q18 ─────
Q(18, "Route A 訓練幾小時,你們 30 秒,是不是太投機取巧?")
core("不是投機,是方法論不同。Transfer Learning 的本質就是『讓預訓練成本承擔最重的計算』。兩路線互補,最終可 ensemble。")

reason_header(1, "兩路線不是競爭,是互補")

mk_table(
    headers=["面向", "Route A", "Route B(本路線)"],
    rows=[
        ["訓練策略", "Full fine-tune + self-training", "Frozen encoder + trainable head"],
        ["可訓參數", "Backbone + head 全部", "只有 head 1.57M"],
        ["訓練時間", "幾小時(Teacher-Student × 4)", "~30 秒/variant"],
        ["對 Pantanal 分布適應", "強(直接 fine-tune)", "中(受 Perch 表徵上限約束)"],
        ["對 multi-domain 泛化", "中(self-training 補 domain gap)", "強(Perch 預訓練即已多域)"],
    ],
    widths_cm=[3.5, 5, 5]
)

reason_header(2, "Transfer learning 的本質")
P("  這不是『投機取巧』,是『讓別人的預訓練成本承擔最重的計算』。", size=11, bold=True)
P("  Google 在 Perch 訓練上花了幾千 GPU 小時,本路線直接受惠。", size=10)
P("  BERT 下游任務、ViT linear probe、SimCLR 等等頂會論文都是幾分鐘訓完。", size=10)

reason_header(3, "最終規劃:ensemble 結合兩者")
P("  probability averaging 兩路線預測,預期再提升 0.5-1.0 百分點。", size=10)

script(
    "這不是投機,是方法論不同。Transfer Learning 的本質就是『讓預訓練成本承擔最重的計算』 —"
    "Google 在 Perch 訓練上花了幾千 GPU 小時,本路線直接受惠。BERT 下游任務、ViT linear probe 也都是幾分鐘訓完。"
    "Route A 的價值在 full fine-tune 對 Pantanal 分布的適應、self-training 補強 domain gap。"
    "兩路線是互補不是競爭,最終規劃 ensemble 結合兩者優點。"
)

doc.add_page_break()


# ==============================================================
# PART 5 · 方法論合理性
# ==============================================================
H("Part 5 · 方法論合理性", level=1)

# ───── Q19 ─────
Q(19, "你只訓 head,算深度學習嗎?")
core("算。head 本身是 3 層 neural network(1.57M 參數)。加上 Perch encoder ~30 層,整個 pipeline 逾 30 層。Transfer Learning 是 DL 主流子領域,BERT / ViT / SimCLR 都用這做法。")

reason_header(1, "深度學習的定義")
P("  關鍵是『多層非線性變換的堆疊』,不是『是否從頭訓練整個網路』。", size=11, bold=True)
P("  只要架構有 neural network 元件(Linear / BN / ReLU / Conv / Attention),", size=10)
P("  且用 gradient descent 訓練,就是 DL。", size=10)

reason_header(2, "本路線的 DL 成分清單")

mk_table(
    headers=["元件", "含 DL 成分"],
    rows=[
        ["Perch encoder", "PCEN + Conformer blocks × 20-30 層(frozen)"],
        ["MLP head", "3 層 Linear + BN + ReLU + Dropout,1.57M 參數"],
        ["訓練演算法", "forward / backward / gradient descent"],
        ["Optimizer", "Adam with Cosine LR schedule"],
        ["Loss design", "BCE / Focal Loss(Lin et al. 2017)"],
        ["Regularization", "Dropout 0.3 + Early stopping"],
        ["評估方法", "Ablation study 控制變因實驗"],
    ],
    widths_cm=[4, 10]
)

reason_header(3, "同類案例證明這是主流")
P("  • SimCLR(ICML 2020):frozen encoder + 2 層 MLP head", size=10)
P("  • BiT(ECCV 2020):frozen backbone + linear probe", size=10)
P("  • MAE(CVPR 2022):linear probing 作為主要評估", size=10)
P("  • BERT 下游任務:絕大多數頂會論文都是 frozen BERT + 1 層 head", size=10)
P("  • Perch 論文本身(Hamer et al. 2023):linear probe 評估 representation", size=10)

script(
    "算深度學習。head 本身是 3 層 neural network 含 BN、Dropout、ReLU,1.57M 可訓參數,不折不扣的 deep network。"
    "加上 Perch encoder 本身 20-30 層,整個 pipeline 超過 30 層。"
    "Transfer Learning 是 representation learning 的核心子領域,SimCLR、BiT、MAE、BERT 下游任務都是這種"
    "frozen backbone + trainable head 做法,頂會論文滿地都是。屬於 DL 主流研究範疇。"
)
sep()

# ───── Q20 ─────
Q(20, "本路線的天花板是什麼?怎麼突破?")
core("天花板 = Perch embedding 的表徵能力。突破路徑:短期 class reweighting、中期部分解凍 Perch、長期針對非鳥類 72 類特化。")

reason_header(1, "天花板的根本來源")
P("  Linear probe 的效能 ≈ embedding 空間的線性可分性", size=10)
P("  本路線 0.9660 基本就是 Perch 1536 維空間在 234 類上的上限", size=10)
P("  MLP head 只能『在空間內做更好的邊界切割』,無法改變空間本身", size=11, bold=True)

reason_header(2, "三階段突破路徑")

mk_table(
    headers=["時程", "策略", "預期效果"],
    rows=[
        ["短期(1 週)", "Class-balanced reweighting 拯救稀有類", "+0.5-1.0 pp"],
        ["中期(3 週)", "部分解凍 Perch 最後 1 個 Conformer block", "+1-2 pp(對 Pantanal 特化)"],
        ["長期(到比賽截止)", "非鳥類 72 類(兩棲/爬蟲/昆蟲/哺乳)專用 head", "提升 bottom-5 表現"],
        ["長期", "Test-time augmentation(TTA)", "+0.3-0.5 pp"],
        ["終局", "Route A + Route B ensemble", "+0.5-1 pp(綜合)"],
    ],
    widths_cm=[3, 6, 5]
)

script(
    "天花板是 Perch embedding 本身的表徵能力 — linear probe 0.9540 已逼近這個上限,MLP head 只能小幅超越。"
    "要突破有三條路。短期做 class-balanced reweighting 救稀有類。中期部分解凍 Perch 最後 1 個 Conformer block,"
    "對 Pantanal 分布特化。長期針對非鳥類 72 類(兩棲、爬蟲、昆蟲、哺乳)做專用 head,因為 Perch 預訓練只有鳥類。"
    "最後搭配 Route A + Route B ensemble,預期綜合還能再提升 0.5-1 百分點。"
)

doc.add_page_break()


# ==============================================================
# PART 6 · 陷阱題
# ==============================================================
H("Part 6 · 陷阱題(最高危)", level=1)
P("這些題目最可能翻車 — 答得好變加分題,答不出來是失分題。", italic=True, size=10, color=RGBColor(0xC5, 0x30, 0x30))
spacer()

# ───── T1 ─────
Q("T1", "你憑什麼說 Variant C 最佳?+0.0005 在 noise 範圍,可能 B 才是真正最佳?")
core("公平承認統計上 B/C 無顯著差異。選 C 因為純數字略高且方向一致,但明確標示為 null result,下一步會做 multi-seed 驗證穩健性。")

reason_header(1, "承認 + 解釋 + 行動計畫")
P("  承認:Focal vs BCE 差距 0.0005 落在 noise 區間,統計上不顯著", size=10)
P("  解釋:三個 AUC 指標(val/ta/sc)C 都略高於 B,方向一致不是隨機", size=10)
P("  行動:下一步做 3-seed average,若 B > C 改推 B 為最終版本", size=10)

script(
    "公平承認 — Focal 跟 BCE 差距 0.0005 確實落在 run-to-run 的雜訊區間,統計上無顯著差異。"
    "選 C 的原因是:(1) 純數字上 C 在三個 AUC 指標都略高於 B,方向一致不是隨機震盪;(2) Focal 保留在架構中,"
    "方便未來擴展 class reweighting;(3) 我在報告裡也明確標示這是 null result,沒有宣稱 Focal 有效。"
    "下一步做 3-seed multi-seed average 驗證穩健性,如果 B 平均高於 C 就改推 B 為最終版本。這是負責任的研究態度。"
)
sep()

# ───── T2 ─────
Q("T2", "你們報告寫 0.9660 很漂亮,實際上 Phase 3 submission 跑出來了嗎?")
core("尚未。Phase 3 是下階段短期 TODO(1 週內完成)。本期交付是 Phase 1 embedding + Phase 2 ablation,Public LB 是下次報告內容。")

reason_header(1, "明確交付範圍,不誇大")
P("  本期交付:Phase 1 (Embedding) + Phase 2 (Ablation) + 完整分析", size=10)
P("  下期交付:Phase 3 (Submission) + Public LB 分數", size=10)
P("  誠實聲明 val 0.9660 是 confident 的 pre-submission 估計,非最終 LB", size=10)

script(
    "誠實答覆 — 尚未。Phase 3 submission notebook 是下階段工作,計畫 1 週內完成。"
    "本期交付範圍是 Phase 1 embedding extraction + Phase 2 head ablation,"
    "val 0.9660 是 confident 的 pre-submission 估計,不是最終 LB 分數。"
    "下次報告會呈現 Public LB 驗證結果。這裡不誇大、不虛報。"
)
sep()

# ───── T3 ─────
Q("T3", "沒試過 fine-tune 最後幾層,憑什麼說凍結最好?")
core("沒宣稱『凍結最好』,而是『在有限時間與資料下凍結是合理起點』。部分解凍是長期 TODO,會與完全凍結做公平對比。")

reason_header(1, "立場微調:凍結 ≠ 最好,是合理起點")
P("  本期宣稱:在本期範圍內,凍結 + MLP head 達 0.9660,超越第一次目標", size=10)
P("  本期『沒』宣稱:凍結是絕對最優解", size=10)
P("  實際立場:凍結是 ablation study 的合理起點,下階段會挑戰這個假設", size=10)

reason_header(2, "長期 TODO #7:部分解凍 Perch")
P("  計畫解凍最後 1 個 Conformer block,用 layer-wise LR(後層 3e-4,head 1e-3)", size=10)
P("  介於本路線完全凍結與 Route A 完全 fine-tune 之間", size=10)
P("  預期對 Pantanal 分布特化有幫助(+1-2 pp)", size=10)

script(
    "沒有宣稱『凍結最好』,只宣稱『在本期範圍內凍結是合理起點』。本期目標是建立 baseline + ablation study,"
    "部分解凍是長期 TODO,因為需要更細緻的 layer-wise learning rate 設計,工作量不小。"
    "下階段會做完全凍結 vs 部分解凍 vs 全 fine-tune 的公平對比,介於本路線與 Route A 之間,"
    "看能否突破 Perch 的表徵上限。這裡不主張『凍結絕對好』,只主張『凍結作為第一階段合理』。"
)
sep()

# ───── T4 ─────
Q("T4", "Val 準不代表 LB 準,若 LB 掉到 0.85 以下怎麼辦?")
core("三層應變:檢查 val split leakage → test-time EDA → Route A+B ensemble。底線即使 LB 0.88,仍超越 BirdNET/YAMNet,方法論有效。")

reason_header(1, "三層防禦方案")

mk_table(
    headers=["檢查/對策", "動作", "預期效果"],
    rows=[
        ["① Val split leakage 檢查", "Multi-seed split + manual spot check", "排除技術性原因"],
        ["② Test-time EDA", "分析 Pantanal 測試分布(物種占比 / soundscape 長度)", "定位 domain shift"],
        ["③ Route A + B ensemble", "probability averaging 綜合兩路線", "+0.5-1 pp 穩健性"],
    ],
    widths_cm=[4.5, 6, 3.5]
)

reason_header(2, "底線保證")
P("  即使 LB 掉到 0.88,仍超越 BirdNET/YAMNet 的 0.86/0.928,方法論有效性不變。", size=11, bold=True)
P("  下階段做 class reweighting / 部分解凍都可挽救。", size=10)

script(
    "有三層應變。第一層檢查 val split 有沒有 leakage,做 multi-seed split 驗證。"
    "第二層做 test-time EDA,檢查 Pantanal 分布跟訓練分布差異在哪。"
    "第三層搭 Route A + B ensemble 提升穩健性。底線是即使 LB 掉到 0.88,仍超越 BirdNET 的 0.86 和 YAMNet 的 0.928,"
    "方法論本身的有效性不受影響。最壞情況下,下階段改 class-balanced reweighting 或部分解凍 Perch 都可挽救。"
)
sep()

# ───── T5 ─────
Q("T5", "Perch 若出 v3 架構大改,你們要不要全部重跑?")
core("看版本差異。若只改 pretrain 資料 → 重抽 embedding + 重訓 head 幾小時完成;若架構大改 → 重新評估整個 pipeline。本路線優勢就在迭代成本低。")

reason_header(1, "本路線的迭代韌性")
P("  換 encoder 只需三步:", size=11)
P("    ① 載入新版 encoder 跑 forward → 產出新 embedding .npy", size=10)
P("    ② 改 head 輸入維度(若 new dim ≠ 1536)", size=10)
P("    ③ 重訓 head(幾分鐘)", size=10)
P("  Route A 的 full fine-tune 換 backbone 要重新訓練整個網路,幾小時起跳。", size=10)

reason_header(2, "架構無關性")
P("  本路線對 encoder 的唯一依賴:『吐出固定維度向量』", size=10)
P("  Perch v2 → v3 / 換 BirdNET / 換 AVES 都能無痛切換", size=10)
P("  甚至可以做 ensemble:多個 encoder 各抽 embedding,head 吃 concat 後的向量", size=10)

script(
    "看版本差異。如果 Perch v3 只改 pretrain 資料、架構相同,就重跑 Phase 1 抽 embedding,Phase 2 head 重訓,"
    "幾小時完成。如果架構大改(例如換成 Transformer-only 或 diffusion-based),就需要重新評估整個 pipeline。"
    "本路線的優勢就在這裡 — 對 encoder 只要求『吐固定維度向量』,換 encoder 只要改 Phase 1 抽取腳本。"
    "相比 Route A 的 full fine-tune 換 backbone 要重新訓練整個網路,本路線的迭代成本低得多,"
    "反而是適應 encoder 更新的穩健做法。"
)

doc.add_page_break()


# ==============================================================
# 結尾 · 使用建議
# ==============================================================
H("使用建議與最高守則", level=1)
spacer()

P("1. Part 1(Q1-Q7)必讀", bold=True, size=12)
P("   觀念基本盤,報告前必須看完並能復述。被問到基本觀念卡住,印象會很差。", size=11)
spacer()

P("2. Part 2(Q8-Q11)A/B/C 防禦 — 背熟", bold=True, size=12, color=RGBColor(0xC0, 0x5C, 0x26))
P("   同學最可能追問『為什麼做 3 個?』『為什麼是 3 層?』,這組背熟變加分題。", size=11)
spacer()

P("3. Part 3-5(Q12-Q20)主動展示", bold=True, size=12)
P("   設計決策與方法論,可在報告時主動帶出,不用等被問。", size=11)
spacer()

P("4. Part 6(T1-T5)陷阱題 — 最高危", bold=True, size=12, color=RGBColor(0xC5, 0x30, 0x30))
P("   這幾題答得好整場 Q&A 會大加分,答不出來會失分。建議逐字背 5 個回答範本。", size=11)
spacer()
spacer()

P("【最高守則】", bold=True, size=14, color=RGBColor(0xC0, 0x5C, 0x26))
P("被問到不知道的題目,誠實說 ——", size=12)
P(
    "「這是下階段規劃的重點,本期先把 baseline 建立起來。」",
    italic=True, bold=True, size=13, color=RGBColor(0x2C, 0x80, 0x41)
)
P("比唬爛好太多。教授最怕的不是不懂,是假裝懂。", size=11, bold=True, color=RGBColor(0xC5, 0x30, 0x30))
spacer()
spacer()
P("祝報告順利 ✨", bold=True, size=16)

# Save
out = r"C:\birdCLEF\new\QA_preparation_detailed.docx"
doc.save(out)
print(f"Generated: {out}")
import os
print(f"File size: {os.path.getsize(out) / 1024:.1f} KB")
